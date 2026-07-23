from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "HUONG_DAN_CAU_HINH_CLOUDFLARE_TAI_SU_DUNG.docx"

BLUE = "1F4D78"
BLUE_LIGHT = "E8EEF5"
INK = "172B3A"
MUTED = "5C6770"
GREEN = "287D3C"
GOLD = "8A6500"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_run_font(a, bold=True)
        b = p.add_run(text[len(bold_prefix):])
        set_run_font(b)
    else:
        set_run_font(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    set_run_font(p.add_run(text))
    return p


def add_numbered_steps(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p_pr.append(num_pr)
        set_run_font(p.add_run(text))


def add_check(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("☐ "), color=BLUE, bold=True)
    set_run_font(p.add_run(text))
    return p


def add_callout(doc, label, text, kind="info"):
    fills = {"info": BLUE_LIGHT, "warn": "FFF4CE", "risk": "FDE8E8", "ok": "E7F4EA"}
    colors = {"info": BLUE, "warn": GOLD, "risk": RED, "ok": GREEN}
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fills[kind])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}: "), color=colors[kind], bold=True)
    set_run_font(p.add_run(text), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        shade(hdr.cells[i], BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=10, color=WHITE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            set_run_font(p.add_run(text), size=10)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.2

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(header.add_run("CLOUDFLARE DEPLOYMENT GUIDE  •  REUSABLE SOP"), size=8.5, color=MUTED, bold=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(footer.add_run("Tài liệu vận hành mẫu — cập nhật 23/07/2026"), size=8.5, color=MUTED)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(48)
p.paragraph_format.space_after = Pt(10)
set_run_font(p.add_run("HƯỚNG DẪN CẤU HÌNH CLOUDFLARE"), size=25, color=BLUE, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(22)
set_run_font(p.add_run("Mẫu triển khai tái sử dụng cho domain và subdomain"), size=15, color=MUTED)

add_callout(
    doc,
    "Kiến trúc mục tiêu",
    "Domain/Registrar → Cloudflare Free (DNS, TLS, CDN, WAF) → Origin hiện có (Vercel, Netlify hoặc hosting khác). Không cần tạo Pages hay Workers nếu Cloudflare chỉ đứng trước origin.",
    "info",
)
add_body(doc, "Đối tượng sử dụng: quản trị viên domain, người vận hành website và người chịu trách nhiệm rollback.")
add_body(doc, "Phạm vi: onboarding zone, chuyển nameserver, cấu hình proxy, TLS, cache, WAF cơ bản, kiểm thử, giám sát và rollback.")
add_body(doc, "Không bao gồm: migration framework, chuyển hosting sang Cloudflare Pages/Workers, hay thay đổi backend nghiệp vụ.")

add_heading(doc, "Cách dùng tài liệu", 1)
add_bullet(doc, "Thay các chuỗi {DOMAIN}, {HOSTNAME}, {ORIGIN_TARGET} bằng giá trị của dự án.")
add_bullet(doc, "Thực hiện theo thứ tự; không bật Proxy trước khi DNS và origin đã được kiểm tra.")
add_bullet(doc, "Lưu bằng chứng trước/sau cho mỗi thay đổi: ảnh chụp, export DNS, thời gian và người thực hiện.")
add_callout(doc, "Điểm dừng bắt buộc", "Nếu thiếu DNS export, quyền registrar, quyền Cloudflare hoặc đường rollback về origin, không thay nameserver.", "risk")

doc.add_page_break()

add_heading(doc, "1. Thu thập thông tin trước triển khai", 1)
add_check(doc, "Domain gốc: {DOMAIN}")
add_check(doc, "Hostname cần proxy: {HOSTNAME}.{DOMAIN}")
add_check(doc, "Origin target do Vercel/Netlify/hosting cung cấp: {ORIGIN_TARGET}")
add_check(doc, "Tài khoản registrar và người có quyền thay nameserver")
add_check(doc, "Tài khoản Cloudflare có quyền quản lý zone")
add_check(doc, "Tài khoản origin và URL mặc định để kiểm tra/rollback")
add_check(doc, "Danh sách webhook, API, admin path và dịch vụ email")
add_check(doc, "Network log/proxy/firewall evidence theo chính sách tổ chức")

add_heading(doc, "2. Sao lưu DNS hiện tại", 1)
add_body(doc, "Xuất hoặc chụp toàn bộ DNS tại nhà cung cấp hiện tại. Phải rà đủ các loại record sau:")
for item in ("A, AAAA và CNAME", "MX", "TXT (SPF, DKIM, DMARC, xác minh dịch vụ)", "CAA", "SRV và record không phục vụ HTTP", "Record wildcard nếu có"):
    add_bullet(doc, item)
add_callout(doc, "Không xóa DNS cũ", "Sau khi đổi nameserver, dữ liệu tại nhà cung cấp cũ không còn authoritative nhưng vẫn hữu ích cho rollback và đối chiếu.", "warn")

add_heading(doc, "3. Quyết định Proxy status", 1)
add_simple_table(
    doc,
    ["Loại record", "Trạng thái ban đầu", "Sau nghiệm thu"],
    [
        ("Hostname web cần Cloudflare bảo vệ", "DNS only", "Proxied"),
        ("Hostname web khác chưa kiểm thử", "DNS only", "Giữ DNS only"),
        ("MX / SPF / DKIM / DMARC", "DNS only", "DNS only"),
        ("TXT xác minh Vercel/Netlify", "DNS only", "DNS only"),
        ("Dịch vụ không phải HTTP/HTTPS", "DNS only", "DNS only"),
    ],
    [3300, 2500, 3560],
)

add_heading(doc, "4. Onboard zone vào Cloudflare", 1)
add_numbered_steps(doc, (
    "Trong Cloudflare chọn Add a domain / Onboard a domain, nhập {DOMAIN}, chọn Free plan.",
    "Cho phép Cloudflare import DNS tự động, sau đó đối chiếu từng record với bản export.",
    "Đặt toàn bộ CNAME web ở DNS only trong giai đoạn đầu.",
    "Kiểm tra DNSSEC/DS tại registrar. Nếu DNSSEC đang bật, xử lý theo hướng dẫn của registrar trước khi đổi nameserver.",
    "Bấm Continue to activation và ghi lại hai nameserver Cloudflare được cấp.",
))

add_heading(doc, "5. Đổi nameserver tại registrar", 1)
add_numbered_steps(doc, (
    "Mở khu vực quản lý tên miền, chọn Cài đặt NS/Nameserver — không phải màn hình chỉnh DNS record.",
    "Thay toàn bộ nameserver cũ bằng đúng hai nameserver Cloudflare.",
    "Không trộn nameserver cũ và Cloudflare; không lặp một nameserver vào nhiều ô.",
    "Lưu thay đổi và quay lại Cloudflare yêu cầu kiểm tra.",
    "Chờ zone chuyển Active; thời gian thường 1–2 giờ nhưng có thể tới 24 giờ.",
))
add_callout(doc, "Trong lúc chờ", "Không bật Proxy, không bật DNSSEC và không thay đổi nameserver liên tục.", "warn")

add_heading(doc, "6. Kiểm tra khi zone đã Active", 1)
for item in (
    "Mọi hostname/subdomain cũ vẫn truy cập được ở chế độ DNS only.",
    "Email gửi/nhận bình thường nếu domain có email.",
    "Origin dashboard vẫn xác nhận custom domain hợp lệ.",
    "Chứng chỉ HTTPS tại origin còn hiệu lực.",
    "Không có record bị thiếu so với DNS export.",
):
    add_check(doc, item)

add_heading(doc, "7. Cấu hình SSL/TLS", 1)
add_simple_table(
    doc,
    ["Thiết lập", "Giá trị khuyến nghị"],
    [
        ("SSL/TLS encryption mode", "Full (strict)"),
        ("Always Use HTTPS", "On"),
        ("Minimum TLS Version", "TLS 1.2"),
        ("TLS 1.3", "On"),
        ("Automatic HTTPS Rewrites", "On"),
        ("Universal SSL", "On"),
        ("HSTS", "Off trong 48–72 giờ đầu"),
    ],
    [4200, 5160],
)
add_callout(doc, "Không dùng Flexible SSL", "Flexible chỉ mã hóa từ browser tới Cloudflare và dễ gây redirect loop với origin HTTPS.", "risk")

add_heading(doc, "8. Cache Rules tối thiểu", 1)
add_body(doc, "Tạo rule bypass trước khi bật Proxy. Ví dụ cho hostname {HOSTNAME}.{DOMAIN}:")
add_bullet(doc, "API: URI Path starts with /api/ → Bypass cache")
add_bullet(doc, "Admin: URI Path starts with /admin/ → Bypass cache")
add_bullet(doc, "Webhook đặc thù: thêm path riêng nếu không nằm dưới /api/")
add_body(doc, "Giữ cache mặc định cho phần còn lại. Không bật Cache Everything cho toàn hostname khi HTML hoặc nội dung quản trị thay đổi thường xuyên.")

add_heading(doc, "9. Bật Proxy có kiểm soát", 1)
add_numbered_steps(doc, (
    "Trong DNS → Records, chỉ chuyển record {HOSTNAME} sang Proxied (đám mây cam).",
    "Giữ các subdomain chưa kiểm thử ở DNS only.",
    "Chờ 2–5 phút và mở trang bằng cửa sổ ẩn danh.",
    "Kiểm tra response, redirect, ảnh, font, API, admin và webhook.",
))
add_callout(doc, "Rollback nhanh", "Nếu xuất hiện lỗi 525, 526, redirect loop hoặc API bị chặn, chuyển riêng {HOSTNAME} về DNS only.", "risk")

add_heading(doc, "10. WAF và bot protection", 1)
add_bullet(doc, "Bật Cloudflare Managed Rules nếu plan hiện tại cung cấp.")
add_bullet(doc, "Không đặt Managed Challenge/JavaScript Challenge trước webhook máy–máy.")
add_bullet(doc, "Chỉ cho phép method cần thiết trên webhook, thường là POST.")
add_bullet(doc, "Rate-limit endpoint report/log và login admin; chừa đủ biên cho webhook retry.")
add_bullet(doc, "Không chặn quốc gia/ASN trước khi có log chứng minh không ảnh hưởng đối tác và khách hàng.")
add_bullet(doc, "Xem Security Events để phát hiện false positive trước khi siết rule.")

add_heading(doc, "11. Checklist nghiệm thu", 1)
for item in (
    "Trang chủ, trang nội dung và asset tải bình thường trên desktop/mobile.",
    "Admin đăng nhập, đọc, ghi và upload hoạt động.",
    "API động trả no-store và không có CF-Cache-Status: HIT.",
    "Webhook hợp lệ thành công; webhook sai chữ ký bị từ chối.",
    "Form đặt lịch/thanh toán, email và calendar hoạt động.",
    "Không có redirect loop, mixed content, lỗi 525/526 hoặc CORS mới.",
    "Security headers từ origin vẫn xuất hiện qua Cloudflare.",
    "Analytics và Security Events không có false positive nghiêm trọng.",
):
    add_check(doc, item)

add_heading(doc, "12. Kiểm thử bằng lệnh", 1)
add_body(doc, "Thay hostname trước khi chạy. Không đưa token, cookie hoặc secret vào terminal history/log.")
for command in (
    "curl -I https://{HOSTNAME}.{DOMAIN}/",
    "curl -I https://{HOSTNAME}.{DOMAIN}/admin/",
    "curl -i -X POST https://{HOSTNAME}.{DOMAIN}/api/csp-report -H 'Content-Type: application/csp-report' --data '{}'",
):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    shade_table = doc.add_table(rows=1, cols=1)
    set_table_geometry(shade_table, [9000])
    cell = shade_table.cell(0, 0)
    shade(cell, "F2F4F7")
    set_run_font(cell.paragraphs[0].add_run(command), name="Courier New", size=9.5, color=INK)

add_heading(doc, "13. Giám sát 72 giờ đầu", 1)
add_check(doc, "HTTP 4xx/5xx, đặc biệt 401, 403, 429, 502, 525 và 526")
add_check(doc, "Cache status của HTML, static assets, /api/* và /admin/*")
add_check(doc, "Webhook retry, duplicate và latency tới origin")
add_check(doc, "WAF events và bot events")
add_check(doc, "Email, booking, thanh toán hoặc tác vụ nghiệp vụ quan trọng")
add_check(doc, "Egress/network evidence và endpoint ngoài danh sách dự kiến")

add_heading(doc, "14. DNSSEC và HSTS", 1)
add_body(doc, "Chỉ thực hiện sau khi hệ thống ổn định ít nhất 48–72 giờ.")
add_bullet(doc, "DNSSEC: bật tại Cloudflare, sao chép chính xác DS record Cloudflare cung cấp sang registrar, rồi xác minh trạng thái.")
add_bullet(doc, "HSTS: chỉ bật sau khi chắc chắn mọi hostname cần thiết đều hỗ trợ HTTPS lâu dài.")
add_bullet(doc, "Không bật includeSubDomains hoặc preload theo mặc định; sai cấu hình có thể làm nhiều subdomain không truy cập được.")

add_heading(doc, "15. Quy trình rollback", 1)
add_numbered_steps(doc, (
    "Ghi lại thời điểm, lỗi, Cloudflare Ray ID và request ID liên quan.",
    "Chuyển record {HOSTNAME} từ Proxied về DNS only.",
    "Kiểm tra trực tiếp origin hostname và custom domain.",
    "Nếu lỗi do deployment origin, rollback tại Vercel/Netlify/hosting.",
    "Không xóa zone, deployment hoặc log trong ngày sự cố.",
    "Sau khi ổn định, điều tra nguyên nhân trước khi bật Proxy lại.",
))

add_heading(doc, "16. Hồ sơ bàn giao", 1)
for item in (
    "DNS export trước/sau",
    "Hai nameserver Cloudflare được cấp",
    "Ảnh/export TLS, Cache Rules và WAF rules",
    "Origin deployment ID hoặc commit SHA",
    "Kết quả kiểm thử và thời gian nghiệm thu",
    "Network evidence theo chính sách tổ chức",
    "Người chịu trách nhiệm vận hành và rollback",
):
    add_check(doc, item)
add_callout(doc, "Bảo mật", "Không lưu secret, API key, cookie, private key, raw webhook chứa dữ liệu khách hàng hoặc thông tin thanh toán trong tài liệu bàn giao.", "risk")

add_heading(doc, "Phụ lục A — Phiếu thông tin dự án", 1)
add_simple_table(
    doc,
    ["Trường", "Giá trị cần điền"],
    [
        ("Domain", "{DOMAIN}"),
        ("Hostname được proxy", "{HOSTNAME}.{DOMAIN}"),
        ("Origin provider", "Vercel / Netlify / Hosting khác"),
        ("Origin target", "{ORIGIN_TARGET}"),
        ("Registrar", "{REGISTRAR}"),
        ("Cloudflare account/zone owner", "{OWNER}"),
        ("Webhook/API paths", "{PATHS}"),
        ("Rollback owner", "{ROLLBACK_OWNER}"),
        ("Ngày triển khai", "{DATE}"),
    ],
    [3300, 6060],
)

add_heading(doc, "Phụ lục B — Tiêu chí hoàn tất", 1)
add_check(doc, "Zone Active, TLS Full (strict), hostname mục tiêu Proxied")
add_check(doc, "API/admin bypass cache")
add_check(doc, "Nghiệp vụ end-to-end đạt")
add_check(doc, "Có rollback đã thử hoặc được xác minh")
add_check(doc, "Có bằng chứng vận hành và người nhận bàn giao")

doc.save(OUTPUT)
print(OUTPUT)
